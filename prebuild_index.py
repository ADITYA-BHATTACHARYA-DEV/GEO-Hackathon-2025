# # prebuild_index.py
# import os
# import sys
# import time
# import json
# import math
# import hashlib
# import concurrent.futures
# from pathlib import Path
# from typing import List, Dict, Optional
#
# # parsing libs
# import fitz  # PyMuPDF
# from docx import Document as DocxDocument
# import openpyxl
# from PIL import Image
# import pytesseract
#
# # progress
# from tqdm import tqdm
#
# # embeddings & chroma
# from langchain_community.embeddings import OllamaEmbeddings
# import chromadb
# from chromadb.config import Settings
#
# # chunking
# from langchain_text_splitters import RecursiveCharacterTextSplitter
# from langchain_core.documents import Document as LCDocument
#
# # =========================
# # Config — change if needed
# # =========================
# ROOT_DIR = "Training data-shared with participants"  # folder with per-well subfolders
# RESULT_DIR = "VectorDB"  # root folder to store per-well chroma DBs
# EMBED_MODEL = "nomic-embed-text"         # recommended fast embedding in Ollama
# EMBED_BATCH_SIZE = 16                    # embedding batch size
# EMBED_WORKERS = 6                        # parallel threads for embedding
# PARSE_WORKERS = 6                        # threads to parse files concurrently
# CHUNK_SIZE = 2500
# CHUNK_OVERLAP = 150
# MIN_TEXT_LEN_FOR_OCR = 200               # if PDF page text < this, run OCR for that page
# SUPPORTED_EXT = (".pdf", ".docx", ".xlsx", ".txt", ".md", ".png", ".jpg", ".jpeg", ".bmp", ".tiff")
# TESSERACT_LANG = "eng"                   # adjust if needed
# # =========================
#
# # ensure output dir
# Path(RESULT_DIR).mkdir(parents=True, exist_ok=True)
#
# # helper: read file mtimes to detect changed files (optional incremental runs)
# def build_mtime_map(folder: Path) -> Dict[str, float]:
#     m = {}
#     for p in folder.rglob("*"):
#         if p.is_file() and p.suffix.lower() in SUPPORTED_EXT:
#             m[str(p)] = p.stat().st_mtime
#     return m
#
# # small logger
# def log(*args, **kwargs):
#     print(time.strftime("[%H:%M:%S]"), *args, **kwargs)
#     sys.stdout.flush()
#
# # -------------------------
# # Parsers
# # -------------------------
# def parse_pdf_text_and_images(path: str, do_ocr_if_scanned: bool = True) -> str:
#     """
#     Use PyMuPDF to extract textual content page-by-page.
#     If page extraction returns little text, optionally OCR the page via pytesseract.
#     """
#     try:
#         doc = fitz.open(path)
#     except Exception as e:
#         return f"[PDF ERROR] {path}: {e}"
#
#     all_text = []
#     for page in doc:
#         text = page.get_text() or ""
#         if do_ocr_if_scanned and len(text.strip()) < MIN_TEXT_LEN_FOR_OCR:
#             # render the page to an image and OCR it
#             try:
#                 mat = fitz.Matrix(2, 2)  # render at higher resolution for better OCR
#                 pix = page.get_pixmap(matrix=mat, alpha=False)
#                 img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
#                 ocr_text = pytesseract.image_to_string(img, lang=TESSERACT_LANG)
#                 text = (text + "\n" + ocr_text).strip()
#             except Exception as e:
#                 # fallback: keep extracted text (even if small)
#                 text = text
#         all_text.append(text)
#     return "\n\n".join(all_text)
#
# def parse_docx(path: str) -> str:
#     try:
#         doc = DocxDocument(path)
#         return "\n".join([p.text for p in doc.paragraphs])
#     except Exception as e:
#         return f"[DOCX ERROR] {path}: {e}"
#
# def parse_xlsx(path: str) -> str:
#     try:
#         wb = openpyxl.load_workbook(path, data_only=True)
#         sheets_text = []
#         for sheet in wb.sheetnames:
#             ws = wb[sheet]
#             rows = []
#             for row in ws.iter_rows(values_only=True):
#                 rows.append("| " + " | ".join("" if v is None else str(v) for v in row) + " |")
#             if rows:
#                 sheets_text.append(f"## Sheet: {sheet}\n" + "\n".join(rows))
#         return "\n\n".join(sheets_text)
#     except Exception as e:
#         return f"[XLSX ERROR] {path}: {e}"
#
# def parse_textfile(path: str) -> str:
#     try:
#         return Path(path).read_text(encoding="utf-8", errors="ignore")
#     except Exception:
#         return Path(path).read_text(encoding="latin-1", errors="ignore")
#
# def parse_image_marker(path: str) -> str:
#     return f"[IMAGE] {path} (no OCR applied by default)"
#
# def parse_file(path: str) -> str:
#     ext = Path(path).suffix.lower()
#     if ext == ".pdf":
#         return parse_pdf_text_and_images(path, do_ocr_if_scanned=True)
#     elif ext == ".docx":
#         return parse_docx(path)
#     elif ext in (".xlsx", ".xls"):
#         return parse_xlsx(path)
#     elif ext in (".txt", ".md"):
#         return parse_textfile(path)
#     elif ext in (".png", ".jpg", ".jpeg", ".bmp", ".tiff"):
#         # For images, we do OCR too (optional): we try OCR here
#         try:
#             img = Image.open(path)
#             return pytesseract.image_to_string(img, lang=TESSERACT_LANG)
#         except Exception:
#             return parse_image_marker(path)
#     else:
#         return ""
#
# # -------------------------
# # Embedding utils
# # -------------------------
# def chunk_documents_for_embedding(docs: List[LCDocument], chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP):
#     splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
#     return splitter.split_documents(docs)
#
# def hash_text(t: str) -> str:
#     return hashlib.sha256(t.encode("utf-8")).hexdigest()
#
# def batchify(lst, n):
#     for i in range(0, len(lst), n):
#         yield lst[i:i+n]
#
# def parallel_batched_embed(embedding_obj, texts: List[str], batch_size=EMBED_BATCH_SIZE, max_workers=EMBED_WORKERS):
#     """
#     Embed a list of texts using batch calls in parallel threads.
#     embedding_obj must provide embed_documents(list_of_texts) -> list[vectors].
#     Returns list[vectors] in same order as texts.
#     """
#     batches = [texts[i:i+batch_size] for i in range(0, len(texts), batch_size)]
#     results = [None] * len(batches)
#
#     def _call(batch_index, batch_texts):
#         return batch_index, embedding_obj.embed_documents(batch_texts)
#
#     with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as exe:
#         futures = [exe.submit(_call, i, b) for i, b in enumerate(batches)]
#         for fut in concurrent.futures.as_completed(futures):
#             i, emb = fut.result()
#             results[i] = emb
#
#     # flatten
#     flattened = []
#     for r in results:
#         flattened.extend(r)
#     return flattened
#
# # -------------------------
# # Chroma client helpers
# # -------------------------
# import chromadb
#
# def chroma_client_for_dir(dirpath: str) -> chromadb.Client:
#     return chromadb.Client(
#         settings=chromadb.config.Settings(
#             chroma_db_impl="duckdb+parquet",
#             persist_directory=dirpath
#         )
#     )
#
#
# # -------------------------
# # Build index for single well
# # -------------------------
# def build_index_for_well(well_path: Path, embedding_obj, show_progress=True):
#     well_id = well_path.name
#     log(f"Building index for well: {well_id}")
#
#     # gather files
#     files = [p for p in well_path.rglob("*") if p.is_file() and p.suffix.lower() in SUPPORTED_EXT]
#     if not files:
#         log(f"No supported files found in {well_id}, skipping.")
#         return
#
#     log(f"Found {len(files)} files for {well_id}")
#
#     # parse files in parallel (IO bound)
#     parsed = []
#     with concurrent.futures.ThreadPoolExecutor(max_workers=PARSE_WORKERS) as exe:
#         futures = {exe.submit(parse_file, str(p)): p for p in files}
#         for fut in tqdm(concurrent.futures.as_completed(futures), total=len(futures), desc=f"Parsing {well_id}"):
#             p = futures[fut]
#             try:
#                 parsed_text = fut.result()
#             except Exception as e:
#                 parsed_text = f"[ERROR PARSING] {p}: {e}"
#             parsed.append((str(p), parsed_text))
#
#     # build LCDocuments (one per file)
#     lcdocs = []
#     for src, txt in parsed:
#         if not txt or (isinstance(txt, str) and not txt.strip()):
#             continue
#         md = {"source": src, "well_id": well_id}
#         lcdocs.append(LCDocument(page_content=txt, metadata=md))
#
#     if not lcdocs:
#         log(f"No textual content found for {well_id}, skipping.")
#         return
#
#     # chunk
#     chunks = chunk_documents_for_embedding(lcdocs, chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP)
#     log(f"{len(chunks)} chunks created for {well_id}")
#
#     # prepare texts and metadata, compute hashes for caching
#     texts = [c.page_content for c in chunks]
#     metadatas = [c.metadata for c in chunks]
#     text_hashes = [hash_text(t) for t in texts]
#
#     # load embed cache for well
#     cache_file = Path(RESULT_DIR) / well_id / "embed_cache.json"
#     cache_map = {}
#     if cache_file.exists():
#         try:
#             cache_map = json.loads(cache_file.read_text(encoding="utf-8"))
#         except Exception:
#             cache_map = {}
#
#     # find which texts need embedding
#     to_embed_indices = [i for i, h in enumerate(text_hashes) if h not in cache_map]
#     log(f"{len(to_embed_indices)} out of {len(texts)} chunks need embedding for {well_id}")
#
#     embeddings_for_all = [None] * len(texts)
#     # fill from cache
#     for i, h in enumerate(text_hashes):
#         if h in cache_map:
#             embeddings_for_all[i] = cache_map[h]
#
#     # embed needed texts in batches (parallel)
#     if to_embed_indices:
#         texts_to_embed = [texts[i] for i in to_embed_indices]
#         # parallel batched embed
#         try:
#             new_embeddings = parallel_batched_embed(embedding_obj, texts_to_embed, batch_size=EMBED_BATCH_SIZE, max_workers=EMBED_WORKERS)
#         except Exception as e:
#             log("Embedding failed:", e)
#             raise
#
#         # map them back
#         for idx, emb in zip(to_embed_indices, new_embeddings):
#             embeddings_for_all[idx] = emb
#             cache_map[text_hashes[idx]] = emb
#
#         # write cache back
#         try:
#             Path(RESULT_DIR, well_id).mkdir(parents=True, exist_ok=True)
#             cache_file.write_text(json.dumps(cache_map), encoding="utf-8")
#         except Exception as e:
#             log("Failed to write embed cache:", e)
#
#     # filter out image-marker chunks (we don't embed image-only chunks)
#     final_texts = []
#     final_embeddings = []
#     final_metadatas = []
#     for txt, emb, md in zip(texts, embeddings_for_all, metadatas):
#         if txt.strip().startswith("[IMAGE]"):
#             # skip embedding (but keep metadata maybe in a side index if needed)
#             continue
#         final_texts.append(txt)
#         final_embeddings.append(emb)
#         final_metadatas.append(md)
#
#     if not final_texts:
#         log(f"No embeddable chunks for {well_id}.")
#         return
#
#     # Create chroma client & collection (persist per well)
#     well_store_dir = Path(RESULT_DIR) / well_id
#     client = chroma_client_for_dir(str(well_store_dir))
#     collection_name = f"{well_id}_collection"
#     # delete existing collection if present (to rebuild cleanly)
#     try:
#         if collection_name in [c.name for c in client.list_collections()]:
#             client.delete_collection(collection_name)
#     except Exception:
#         pass
#
#     collection = client.create_collection(name=collection_name)
#
#     # build unique ids
#     ids = [f"{well_id}__{i}" for i in range(len(final_texts))]
#
#     # Upsert directly with embeddings (fast)
#     collection.add(
#         ids=ids,
#         metadatas=final_metadatas,
#         documents=final_texts,
#         embeddings=final_embeddings
#     )
#
#     # persist - chroma with duckdb+parquet persists automatically in the directory we passed
#     log(f"Index created for {well_id} at {well_store_dir}")
#
# # -------------------------
# # Main prebuild loop (per well)
# # -------------------------
# def build_all_wells(root_dir: str = ROOT_DIR):
#     root = Path(root_dir)
#     if not root.exists():
#         raise FileNotFoundError(f"{root_dir} not found")
#
#     well_dirs = [d for d in root.iterdir() if d.is_dir()]
#     if not well_dirs:
#         log("No subfolders (wells) found in root directory.")
#         return
#
#     log(f"Preparing embedding model: {EMBED_MODEL}")
#     emb = OllamaEmbeddings(model=EMBED_MODEL)  # local Ollama must be running and model pulled
#
#     for wd in well_dirs:
#         try:
#             build_index_for_well(wd, emb)
#         except Exception as e:
#             log(f"Failed to build index for {wd.name}: {e}")
#
#     log("All done. Per-well vector DBs stored under", RESULT_DIR)
#
#
# if __name__ == "__main__":
#     start = time.time()
#     build_all_wells(ROOT_DIR)
#     log("Total time:", round(time.time() - start, 1), "seconds")




#!/usr/bin/env python3
"""
Optimized prebuild_index.py
Features:
- CPU-friendly embeddings (MiniLM)
- Parallel file parsing & batch embedding
- OCR for PDFs/images
- Chunking for text, tables, plots
- Cache embeddings to speed up reruns
- Stores per-well Chroma collections
"""

import os, sys, time, json, hashlib, concurrent.futures
from pathlib import Path
from typing import List, Tuple
from tqdm import tqdm

import fitz
from docx import Document as DocxDocument
import openpyxl
from PIL import Image
import pytesseract

from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document as LCDocument

from sentence_transformers import SentenceTransformer
import numpy as np
import torch

import chromadb
from chromadb.config import Settings

# ----------------- CONFIG -----------------
ROOT_DIR = Path("Training data-shared with participants")
RESULT_DIR = Path("vectorDB")
RESULT_DIR.mkdir(exist_ok=True, parents=True)

EMB_MODEL = "sentence-transformers/all-MiniLM-L6-v2"
EMBED_BATCH = 256
PARSE_WORKERS = max(2, (os.cpu_count() or 4) // 2)
CHUNK_SIZE = 3500
CHUNK_OVERLAP = 150
MIN_PAGE_TEXT_LEN_FOR_OCR = 250
SUPPORTED_EXT = (".pdf", ".docx", ".xlsx", ".xls", ".txt", ".md", ".png", ".jpg", ".jpeg", ".bmp", ".tiff")
TESSERACT_LANG = "eng"
CHROMA_IMPL = "duckdb+parquet"
EMB_CACHE_FILENAME = "embed_cache.pkl"
EMB_CHECKPOINT_FILENAME = "emb_checkpoint.json"
MTIME_MAP_FILENAME = ".mtimes.json"

# ----------------- UTILITIES -----------------
def log(*args, **kwargs):
    print(time.strftime("[%H:%M:%S]"), *args, **kwargs)
    sys.stdout.flush()

def hash_text(t: str) -> str:
    return hashlib.sha256(t.encode("utf-8")).hexdigest()

def save_pickle(fp: Path, obj):
    with fp.open("wb") as f: import pickle; pickle.dump(obj, f)

def load_pickle(fp: Path):
    import pickle
    if fp.exists():
        with fp.open("rb") as f:
            return pickle.load(f)
    return {}

def save_json(fp: Path, obj):
    with fp.open("w", encoding="utf-8") as f:
        json.dump(obj, f)

def load_json(fp: Path):
    if fp.exists():
        with fp.open("r", encoding="utf-8") as f:
            return json.load(f)
    return {}

# ----------------- PARSERS -----------------
def parse_pdf_text_with_optional_ocr(path: str) -> str:
    try:
        doc = fitz.open(path)
    except: return f"[PDF ERROR] {path}"
    pages_text = []
    for page in doc:
        text = page.get_text() or ""
        if len(text.strip()) < MIN_PAGE_TEXT_LEN_FOR_OCR:
            try:
                pix = page.get_pixmap(matrix=fitz.Matrix(2,2))
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                ocr_text = pytesseract.image_to_string(img, lang=TESSERACT_LANG)
                text += "\n" + ocr_text
            except: pass
        pages_text.append(text)
    return "\n\n".join(pages_text)

def parse_docx(path: str) -> str:
    try: return "\n".join([p.text for p in DocxDocument(path).paragraphs if p.text])
    except: return f"[DOCX ERROR] {path}"

def parse_xlsx(path: str) -> str:
    try:
        wb = openpyxl.load_workbook(path, data_only=True)
        sheets=[]
        for sheet in wb.sheetnames:
            ws=wb[sheet]; rows=[]
            for r in ws.iter_rows(values_only=True):
                rows.append("| " + " | ".join("" if v is None else str(v) for v in r) + " |")
            if rows: sheets.append(f"### Sheet: {sheet}\n" + "\n".join(rows))
        return "\n\n".join(sheets)
    except: return f"[XLSX ERROR] {path}"

def parse_text_file(path: str) -> str:
    try: return Path(path).read_text(encoding="utf-8", errors="ignore")
    except: return Path(path).read_text(encoding="latin-1", errors="ignore")

def parse_image_ocr(path: str) -> str:
    try:
        img = Image.open(path)
        return pytesseract.image_to_string(img, lang=TESSERACT_LANG)
    except: return f"[IMAGE] {path} (no OCR)"

def parse_file(path: str) -> Tuple[str,str]:
    ext = Path(path).suffix.lower()
    if ext==".pdf": return (path, parse_pdf_text_with_optional_ocr(path))
    elif ext==".docx": return (path, parse_docx(path))
    elif ext in (".xlsx",".xls"): return (path, parse_xlsx(path))
    elif ext in (".txt",".md"): return (path, parse_text_file(path))
    elif ext in (".png",".jpg",".jpeg",".bmp",".tiff"):
        txt = parse_image_ocr(path)
        return (path, txt if txt.strip() else f"[IMAGE] {path} (no OCR)")
    return (path, "")

# ----------------- Chroma -----------------
def chroma_client_for_dir(dirpath: str) -> chromadb.Client:
    Path(dirpath).mkdir(parents=True, exist_ok=True)
    return chromadb.Client(Settings(chroma_db_impl=CHROMA_IMPL, persist_directory=str(dirpath)))

# ----------------- Embedding -----------------
def prepare_embedder(model_name: str):
    torch.set_num_threads(max(1, os.cpu_count() or 1))
    log(f"Loading embedding model {model_name} on CPU...")
    return SentenceTransformer(model_name, device="cpu")

def batch_embed_texts_parallel(embedder, texts: List[str], batch_size=EMBED_BATCH):
    vectors = [None]*len(texts)
    def worker(batch_indices):
        batch_texts = [texts[i] for i in batch_indices]
        emb = embedder.encode(batch_texts, batch_size=len(batch_texts), convert_to_numpy=True, show_progress_bar=False)
        for j,i in enumerate(batch_indices): vectors[i]=emb[j].tolist()
    batches=[list(range(i,min(i+batch_size,len(texts)))) for i in range(0,len(texts),batch_size)]
    with concurrent.futures.ThreadPoolExecutor(max_workers=PARSE_WORKERS) as exe:
        list(tqdm(exe.map(worker,batches), total=len(batches), desc="Embedding batches", ncols=90))
    return vectors

# ----------------- Indexing -----------------
def build_index_for_well(well_path: Path, embedder):
    well_id = well_path.name
    log(f"--- Building index for {well_id} ---")
    files = [p for p in well_path.rglob("*") if p.is_file() and p.suffix.lower() in SUPPORTED_EXT]
    if not files: log(f"No files for {well_id}"); return
    log(f"Parsing {len(files)} files...")
    parsed=[]
    with concurrent.futures.ThreadPoolExecutor(max_workers=PARSE_WORKERS) as exe:
        futures={exe.submit(parse_file,str(p)):p for p in files}
        for fut in tqdm(concurrent.futures.as_completed(futures), total=len(futures), desc=f"Parsing {well_id}", ncols=90):
            p=futures[fut]; parsed.append(fut.result())
    docs=[LCDocument(page_content=t, metadata={"source":s,"well_id":well_id}) for s,t in parsed if t.strip()]
    if not docs: log(f"No text found for {well_id}"); return
    splitter=RecursiveCharacterTextSplitter(chunk_size=CHUNK_SIZE,chunk_overlap=CHUNK_OVERLAP)
    chunks=splitter.split_documents(docs)
    log(f"{len(chunks)} chunks created for {well_id}")
    texts=[c.page_content for c in chunks]; metas=[c.metadata for c in chunks]; hashes=[hash_text(t) for t in texts]
    well_dir=RESULT_DIR/well_id; well_dir.mkdir(parents=True, exist_ok=True)
    emb_cache_file=well_dir/EMB_CACHE_FILENAME
    emb_cache=load_pickle(emb_cache_file)
    emb_cache={k:v for k,v in emb_cache.items()} if emb_cache else {}
    need_indices=[i for i,h in enumerate(hashes) if h not in emb_cache]
    log(f"{len(need_indices)}/{len(texts)} chunks need embeddings")
    if need_indices:
        to_embed=[texts[i] for i in need_indices]
        batch_embs=batch_embed_texts_parallel(embedder,to_embed)
        for idx,v in zip(need_indices,batch_embs): emb_cache[hashes[idx]]=v
        save_pickle(emb_cache_file,emb_cache)
    # build Chroma
    final_texts=[]; final_embeddings=[]; final_metadatas=[]; ids=[]
    for i,(t,m) in enumerate(zip(texts,metas)):
        if t.strip().startswith("[IMAGE]"): continue
        vec=emb_cache.get(hashes[i])
        if not vec: continue
        final_texts.append(t); final_embeddings.append(vec); final_metadatas.append(m); ids.append(f"{well_id}__{i}")
    if not final_texts: log(f"No embeddable chunks for {well_id}"); return
    client=chroma_client_for_dir(RESULT_DIR/well_id)
    coll_name=f"{well_id}_collection"
    try: existing=[c.name for c in client.list_collections()];
    except: existing=[]
    if coll_name in existing: client.delete_collection(coll_name)
    coll=client.create_collection(coll_name)
    coll.add(ids=ids,documents=final_texts,metadatas=final_metadatas,embeddings=final_embeddings)
    log(f"Collection '{coll_name}' created for {well_id}")

def build_all_wells(root_dir: Path=ROOT_DIR):
    well_dirs=[d for d in sorted(root_dir.iterdir()) if d.is_dir()]
    if not well_dirs: log("No wells found"); return
    embedder=prepare_embedder(EMB_MODEL)
    for wd in well_dirs:
        try: build_index_for_well(wd, embedder)
        except Exception as e: log(f"ERROR building {wd.name}: {e}")
    log("All wells processed.")

if __name__=="__main__":
    start=time.time(); log("Starting prebuild index...")
    build_all_wells(ROOT_DIR)
    log("Total time (s):", round(time.time()-start,1))
